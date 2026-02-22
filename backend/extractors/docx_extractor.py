"""
ADIVA - DOCX Extractor

Extracts text and structured data from Microsoft Word (.docx) files.
Supports merged cells detection and table structure analysis.
"""

from pathlib import Path
from typing import Dict, Any, List, Tuple, Optional
from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from extractors.base_extractor import BaseExtractor
from logger import logger, log_extraction, log_error
import time


class DOCXExtractor(BaseExtractor):
    """
    Extracts content from DOCX files using python-docx.
    Enhanced with merged cells detection and table structure analysis.
    """

    def __init__(self):
        """Initialize DOCX extractor"""
        super().__init__()
        self.supported_extensions = {".docx"}

    def can_extract(self, file_path: Path) -> bool:
        """Check if this can extract from the file"""
        return file_path.suffix.lower() in self.supported_extensions

    def _is_cell_merged(self, cell: _Cell) -> Tuple[bool, Optional[int], Optional[int]]:
        """
        Check if a cell is part of a merged region.

        Returns:
            (is_merged, row_span, col_span)
        """
        tc = cell._tc
        tcPr = tc.tcPr

        if tcPr is None:
            return False, None, None

        vMerge = tcPr.find(qn("w:vMerge"))
        hMerge = tcPr.find(qn("w:hMerge"))

        row_span = None
        col_span = None
        is_merged = False

        if vMerge is not None:
            is_merged = True
            val = vMerge.get(qn("w:val"), "continue")
            if val == "restart":
                row_span = self._calculate_row_span(cell)

        if hMerge is not None:
            is_merged = True
            val = hMerge.get(qn("w:val"), "continue")
            if val == "restart":
                col_span = self._calculate_col_span(cell)

        gridSpan = tcPr.find(qn("w:gridSpan"))
        if gridSpan is not None:
            is_merged = True
            col_span = int(gridSpan.get(qn("w:val"), 1))

        return is_merged, row_span, col_span

    def _calculate_row_span(self, cell: _Cell) -> int:
        """Calculate the row span for a vertically merged cell."""
        tc = cell._tc
        tcPr = tc.tcPr

        if tcPr is None:
            return 1

        vMerge = tcPr.find(qn("w:vMerge"))
        if vMerge is not None and vMerge.get(qn("w:val"), "") == "restart":
            return 1

        return 1

    def _calculate_col_span(self, cell: _Cell) -> int:
        """Calculate the column span for a horizontally merged cell."""
        tc = cell._tc
        tcPr = tc.tcPr

        if tcPr is None:
            return 1

        gridSpan = tcPr.find(qn("w:gridSpan"))
        if gridSpan is not None:
            return int(gridSpan.get(qn("w:val"), 1))

        return 1

    def _analyze_table_structure(self, table: Table) -> Dict[str, Any]:
        """
        Analyze table structure for merged cells and nesting.

        Returns:
            Dictionary with merge information
        """
        merge_info = {
            "has_merged_cells": False,
            "merged_regions": [],
            "total_cells": 0,
            "merged_cells": 0,
        }

        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                merge_info["total_cells"] += 1

                is_merged, row_span, col_span = self._is_cell_merged(cell)

                if is_merged:
                    merge_info["has_merged_cells"] = True
                    merge_info["merged_cells"] += 1

                    if row_span or col_span:
                        merge_info["merged_regions"].append(
                            {
                                "row": row_idx,
                                "col": col_idx,
                                "row_span": row_span or 1,
                                "col_span": col_span or 1,
                            }
                        )

        return merge_info

    def _extract_table_with_merges(
        self, table: Table
    ) -> Tuple[List[List[str]], List[Dict]]:
        """
        Extract table content with merge information.

        Returns:
            (rows_data, merge_info_list)
        """
        rows_data = []
        merge_info_list = []

        cell_tracker = set()

        for row_idx, row in enumerate(table.rows):
            row_data = []
            col_offset = 0

            for col_idx, cell in enumerate(row.cells):
                cell_id = id(cell)

                actual_col = col_idx + col_offset

                while (row_idx, actual_col) in cell_tracker:
                    actual_col += 1

                text = cell.text.strip()
                row_data.append(text)

                is_merged, row_span, col_span = self._is_cell_merged(cell)

                if is_merged:
                    actual_col_span = col_span or 1
                    actual_row_span = row_span or 1

                    if row_span or col_span:
                        merge_info_list.append(
                            {
                                "row": row_idx,
                                "col": actual_col,
                                "row_span": actual_row_span,
                                "col_span": actual_col_span,
                                "text": text,
                            }
                        )

                    for r in range(actual_row_span):
                        for c in range(actual_col_span):
                            cell_tracker.add((row_idx + r, actual_col + c))
                else:
                    cell_tracker.add((row_idx, actual_col))

            rows_data.append(row_data)

        return rows_data, merge_info_list

    def _validate_table(self, table_data: Dict[str, Any]) -> bool:
        """Validate table structure."""
        if not table_data:
            return False

        min_rows = 1
        min_cols = 2

        if table_data["num_rows"] < min_rows:
            return False

        if table_data["num_cols"] < min_cols:
            return False

        return True

    def extract_text(self, file_path: Path) -> str:
        """
        Extract text from DOCX

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text
        """
        start_time = time.time()

        try:
            doc = Document(str(file_path))
            full_text = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                full_text.append("\n[TABLE]")
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    full_text.append(row_text)
                full_text.append("[/TABLE]\n")

            result = "\n".join(full_text)

            # Log extraction
            extraction_time = time.time() - start_time
            log_extraction(file_path.name, len(result), extraction_time)

            return result

        except Exception as e:
            log_error("DOCXExtraction", str(e), f"File: {file_path}")
            raise

    def extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract metadata from DOCX

        Args:
            file_path: Path to DOCX file

        Returns:
            Metadata dictionary
        """
        try:
            doc = Document(str(file_path))

            metadata = {
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(doc.tables),
                "num_sections": len(doc.sections),
            }

            # Count merged tables
            merged_tables = 0
            for table in doc.tables:
                merge_info = self._analyze_table_structure(table)
                if merge_info["has_merged_cells"]:
                    merged_tables += 1

            metadata["tables_with_merged_cells"] = merged_tables

            # Try to extract core properties
            try:
                core_props = doc.core_properties
                metadata["author"] = core_props.author or ""
                metadata["title"] = core_props.title or ""
                metadata["created"] = (
                    str(core_props.created) if core_props.created else ""
                )
                metadata["modified"] = (
                    str(core_props.modified) if core_props.modified else ""
                )
            except:
                pass

            return metadata

        except Exception as e:
            log_error("DOCXMetadataExtraction", str(e), f"File: {file_path}")
            return {}

    def extract_tables(self, file_path: Path) -> list:
        """
        Extract tables from DOCX with merged cells detection.

        Args:
            file_path: Path to DOCX file

        Returns:
            List of tables with merge information
        """
        try:
            doc = Document(str(file_path))
            all_tables = []

            for table_num, table in enumerate(doc.tables, 1):
                # Analyze table structure
                structure_info = self._analyze_table_structure(table)

                # Extract table data with merge info
                rows_data, merge_list = self._extract_table_with_merges(table)

                table_data = {
                    "table_num": table_num,
                    "num_rows": len(table.rows),
                    "num_cols": len(table.columns),
                    "rows": rows_data,
                    "data": [],
                    "has_merged_cells": structure_info["has_merged_cells"],
                    "merged_regions": merge_list,
                    "merge_statistics": {
                        "total_cells": structure_info["total_cells"],
                        "merged_cells": structure_info["merged_cells"],
                    },
                }

                # Assume first row is header
                if rows_data:
                    headers = rows_data[0]
                    table_data["headers"] = headers

                    # Convert to list of dictionaries
                    for row in rows_data[1:]:
                        if len(row) == len(headers):
                            row_dict = {headers[i]: row[i] for i in range(len(headers))}
                            table_data["data"].append(row_dict)

                if self._validate_table(table_data):
                    all_tables.append(table_data)

                    if structure_info["has_merged_cells"]:
                        logger.info(
                            f"Table {table_num}: {table_data['num_rows']} rows x {table_data['num_cols']} cols (with {structure_info['merged_cells']} merged cells)"
                        )
                    else:
                        logger.info(
                            f"Table {table_num}: {table_data['num_rows']} rows x {table_data['num_cols']} cols"
                        )

            return all_tables

        except Exception as e:
            log_error("DOCXTableExtraction", str(e), f"File: {file_path}")
            return []
